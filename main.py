import os
import sys
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from ttkthemes import ThemedTk
import database
import pandas as pd
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageTk

class App(ThemedTk):
    def __init__(self):
        super().__init__(theme="radiance")

        self.title("InventarioTUA31 - Sistema de Gestión de Inventario para Equipos de Computo")
        self.geometry("850x700")
        self.protocol("WM_DELETE_WINDOW", self.on_closing)
        try:
            # Carga la imagen original
            logo_original = Image.open("tua31v.png") 
            logo_original = logo_original.resize((250, 200), Image.Resampling.LANCZOS)
            self.logo_image = ImageTk.PhotoImage(logo_original)
        except FileNotFoundError:
            print("Advertencia: No se encontró el archivo 'logo'. El logo no se mostrará.")
            self.logo_image = None
        except Exception as e:
            print(f"Error al cargar el logo: {e}")
            self.logo_image = None
        self.mostrar_pantalla_principal()

    def limpiar_pantalla(self):
        for widget in self.winfo_children():
            widget.destroy()

    def on_closing(self):
        if messagebox.askokcancel("Salir", "¿Estás seguro que deseas salir?"):
            self.destroy()

    def mostrar_pantalla_principal(self):
        self.limpiar_pantalla()
        
        main_frame = ttk.Frame(self, padding="20")
        main_frame.pack(expand=True, fill="both")

        if self.logo_image:
            logo_label = ttk.Label(main_frame, image=self.logo_image)
            logo_label.pack(pady=(0, 10))
        ttk.Label(main_frame, text="Sistema de Gestión de Inventario", font=("Arial", 20, "bold")).pack(pady=10)
        
        ttk.Button(main_frame, text="Ver Inventario Completo", command=self.mostrar_inventario_completo).pack(pady=5, ipadx=5, ipady=2.5)
        ttk.Button(main_frame, text="Agregar Nuevo Equipo", command=lambda: self.mostrar_gestion_equipo()).pack(pady=5, ipadx=5, ipady=2.5)

        recientes_frame = ttk.LabelFrame(main_frame, text="Últimos Movimientos", padding="10")
        recientes_frame.pack(pady=10, fill="x", padx=10)
        movimientos_recientes = database.obtener_movimientos_recientes()
        
        if movimientos_recientes:
            for movimiento in movimientos_recientes:
                info, accion, fecha_completa_str = movimiento
                try:
                    fecha = fecha_completa_str.split(" ")[0]
                except:
                    fecha = fecha_completa_str

                texto = f"Acción: {accion} | Detalles: {info} | Fecha: {fecha}"
                
                color = "green"
                if accion == "Actualizado":
                    color = "blue"
                elif accion == "Eliminado":
                    color = "red"
                
                label = ttk.Label(recientes_frame, text=texto, foreground=color)
                label.pack(anchor="w", pady=2)
        else:
            ttk.Label(recientes_frame, text="No hay movimientos recientes.").pack()

    def mostrar_inventario_completo(self):
        self.limpiar_pantalla()
        
        inv_frame = ttk.Frame(self, padding="10")
        inv_frame.pack(expand=True, fill="both")

        # Frame para los botones superiores
        top_buttons_frame = ttk.Frame(inv_frame)
        top_buttons_frame.pack(fill="x", pady=5, padx=5)

        ttk.Button(top_buttons_frame, text="< Volver a Inicio", command=self.mostrar_pantalla_principal).pack(side="left", anchor="w", padx=5)
        
        # Frame para botones de acción
        action_buttons_frame = ttk.Frame(top_buttons_frame)
        action_buttons_frame.pack(side="right", anchor="e")

        # Botones de Exportar
        ttk.Button(action_buttons_frame, text="Exportar a Excel", command=self.exportar_a_excel).pack(side="left", padx=5)
        ttk.Button(action_buttons_frame, text="Exportar a Word", command=self.exportar_a_word).pack(side="left", padx=5)

        # Botón para eliminar
        ttk.Button(action_buttons_frame, text="Eliminar Seleccionado", command=self.eliminar_equipo_seleccionado).pack(side="left", padx=5)


        ttk.Label(inv_frame, text="Inventario Completo de Equipos", font=("Arial", 18, "bold")).pack(pady=10)

        tabla_container = ttk.Frame(inv_frame)
        tabla_container.pack(expand=True, fill="both")

        v_scrollbar = ttk.Scrollbar(tabla_container, orient="vertical")
        h_scrollbar = ttk.Scrollbar(tabla_container, orient="horizontal")

        self.columnas_inventario = ("ID", "Asignado a", "Tipo", "N° Inventario", "Marca", "Modelo", "N° Serie", "Estado")
        
        self.tree = ttk.Treeview(
            tabla_container, 
            columns=self.columnas_inventario, 
            show='headings', 
            yscrollcommand=v_scrollbar.set, 
            xscrollcommand=h_scrollbar.set
        )
        
        v_scrollbar.config(command=self.tree.yview)
        h_scrollbar.config(command=self.tree.xview)

        for col in self.columnas_inventario:
            self.tree.heading(col, text=col)
            self.tree.column(col, width=150, minwidth=100, anchor="w")

        self.tree.column("ID", width=50, minwidth=40, stretch=tk.NO) 
        self.tree.column("N° Serie", width=180)
        v_scrollbar.pack(side="right", fill="y")
        h_scrollbar.pack(side="bottom", fill="x")
        self.tree.pack(side="left", expand=True, fill="both")
        
        self.cargar_datos_inventario()

        self.tree.bind("<Double-1>", self.editar_equipo_seleccionado)

    def cargar_datos_inventario(self):
        for i in self.tree.get_children():
            self.tree.delete(i)
        for equipo in database.obtener_todos_los_equipos():
            self.tree.insert("", "end", values=equipo)

    def editar_equipo_seleccionado(self, event):
        item_id = self.tree.focus()
        if item_id:
            item_values = self.tree.item(item_id, "values")
            id_equipo = item_values[0]
            self.mostrar_gestion_equipo(id_equipo)

    def eliminar_equipo_seleccionado(self):
        item_id_seleccionado = self.tree.focus()
        if not item_id_seleccionado:
            messagebox.showwarning("Sin selección", "Por favor, selecciona un equipo de la lista para eliminar.")
            return

        valores = self.tree.item(item_id_seleccionado, "values")
        id_equipo = valores[0]
        mensaje_confirmacion = f"¿Estás seguro de que deseas eliminar permanentemente el equipo:\n\nID: {id_equipo}\nAsignado a: {valores[1]}\nTipo: {valores[2]}\nN° Inventario: {valores[3]}\n\nEsta acción no se puede deshacer."
        if messagebox.askyesno("Confirmar Eliminación", mensaje_confirmacion):
            if database.eliminar_equipo(id_equipo):
                messagebox.showinfo("Éxito", "El equipo ha sido eliminado correctamente.")
                self.cargar_datos_inventario() 
            else:
                messagebox.showerror("Error", "No se pudo eliminar el equipo de la base de datos.")

    #Exportación
    def exportar_a_excel(self):
        """Exporta los datos del inventario a un archivo Excel (.xlsx)."""
        datos = database.obtener_todos_los_equipos()
        if not datos:
            messagebox.showinfo("Información", "No hay datos en el inventario para exportar.")
            return

        try:
            filepath = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Archivos Excel", "*.xlsx"), ("Todos los archivos", "*.*")],
                title="Guardar inventario como Excel"
            )

            if not filepath:
                return

            df = pd.DataFrame(datos, columns=self.columnas_inventario)

            df.to_excel(filepath, index=False, sheet_name='Inventario TUA31')

            messagebox.showinfo("Éxito", f"Inventario exportado correctamente a:\n{filepath}")

        except Exception as e:
            messagebox.showerror("Error de Exportación", f"Ocurrió un error al exportar a Excel:\n{e}")

    def exportar_a_word(self):
        """Exporta los datos del inventario a un archivo Word (.docx)."""
        datos = database.obtener_todos_los_equipos()
        if not datos:
            messagebox.showinfo("Información", "No hay datos en el inventario para exportar.")
            return

        try:
            filepath = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Archivos Word", "*.docx"), ("Todos los archivos", "*.*")],
                title="Guardar inventario como Word"
            )

            if not filepath:
                return

            # Crear un documento de Word
            doc = Document()
            doc.add_heading('Inventario Completo de Equipos - TUA31', level=1)
            doc.add_paragraph(f"Fecha de exportación: {pd.Timestamp.now().strftime('%d/%m/%Y %H:%M:%S')}") # Usar pandas para la fecha actual
            doc.add_paragraph()

            # Añadir una tabla
            num_cols = len(self.columnas_inventario)
            num_rows = len(datos) + 1 # +1 para la fila de encabezado
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid' # Aplicar un estilo básico

            # Añadir encabezados a la tabla
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(self.columnas_inventario):
                hdr_cells[i].text = col_name
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True

            # Añadir los datos a la tabla
            for idx_row, row_data in enumerate(datos):
                row_cells = table.rows[idx_row + 1].cells
                for idx_col, cell_data in enumerate(row_data):
                    row_cells[idx_col].text = str(cell_data)

            # Opcional: Ajustar anchos de columna (puede requerir más ajustes)
            # for i, width in enumerate([0.5, 1.5, 1.0, 1.0, 1.0, 1.0, 1.5, 0.8]): # Ejemplo de anchos en pulgadas
            #     for cell in table.columns[i].cells:
            #         cell.width = Inches(width)

            # Guardar el documento
            doc.save(filepath)

            messagebox.showinfo("Éxito", f"Inventario exportado correctamente a:\n{filepath}")

        except Exception as e:
            messagebox.showerror("Error de Exportación", f"Ocurrió un error al exportar a Word:\n{e}")
            


    def mostrar_gestion_equipo(self, id_equipo=None):
        self.limpiar_pantalla()
        
        gestion_frame = ttk.Frame(self, padding="20")
        gestion_frame.pack(expand=True)
        
        titulo = "Agregar Nuevo Equipo" if id_equipo is None else f"Actualizar Equipo (ID: {id_equipo})" # Mostrar ID al editar
        ttk.Label(gestion_frame, text=titulo, font=("Arial", 18, "bold")).grid(row=0, column=0, columnspan=2, pady=20)
        
        # Formulario
        labels = ["Asignado a:", "Tipo de Equipo:", "N° Inventario/Arrendado:", "Marca:", "Modelo:", "N° Serie:", "Estado:", "Descripción del Estado:"]
        self.entries = {}
        self.entry_keys = {
            "Asignado a:": "asignado_a",
            "Tipo de Equipo:": "tipo_equipo",
            "N° Inventario/Arrendado:": "num_inventario",
            "Marca:": "marca",
            "Modelo:": "modelo",
            "N° Serie:": "num_serie",
            "Estado:": "estado",
            "Descripción del Estado:": "descripcion_estado"
        }
        
        tipos_equipo = ["PC", "Monitor", "No-Break", "Bocinas", "Camara" , "Laptop", "Impresora", "Disco duro","Camara", "Telefono", "Switch", "Otro"]
        estados_equipo = ["Funcionando", "Sin funcionar", "En diagnóstico", "En reparación", "Para baja", "Baja", "Otro"]

        for i, label_text in enumerate(labels):
            key = self.entry_keys[label_text]
            ttk.Label(gestion_frame, text=label_text).grid(row=i+1, column=0, sticky="w", padx=10, pady=5)
            
            if key == "tipo_equipo":
                self.entries[key] = ttk.Combobox(gestion_frame, values=tipos_equipo, width=38)
            elif key == "estado":
                self.entries[key] = ttk.Combobox(gestion_frame, values=estados_equipo, width=38)
            elif key == "descripcion_estado": 
                text_frame = ttk.Frame(gestion_frame)
                text_scrollbar = ttk.Scrollbar(text_frame)
                self.entries[key] = tk.Text(text_frame, width=40, height=4, yscrollcommand=text_scrollbar.set)
                text_scrollbar.config(command=self.entries[key].yview)
                
                self.entries[key].pack(side="left", fill="both", expand=True)
                text_scrollbar.pack(side="right", fill="y")
                text_frame.grid(row=i+1, column=1, padx=10, pady=5, sticky="ew")
            else:
                self.entries[key] = ttk.Entry(gestion_frame, width=40)

            if key != "descripcion_estado":
                self.entries[key].grid(row=i+1, column=1, padx=10, pady=5, sticky="ew")

        # Cargar datos si es una edición
        if id_equipo:
            self.cargar_datos_equipo(id_equipo)

        # Frame para botones Guardar/Cancelar
        frame_botones = ttk.Frame(gestion_frame)
        frame_botones.grid(row=len(labels)+1, column=0, columnspan=2, pady=20)
        
        btn_guardar = ttk.Button(frame_botones, text="Guardar", command=lambda: self.guardar_equipo(id_equipo))
        btn_guardar.pack(side="left", padx=10)
        
        comando_cancelar = self.mostrar_inventario_completo if id_equipo else self.mostrar_pantalla_principal
        btn_cancelar = ttk.Button(frame_botones, text="Cancelar", command=comando_cancelar)
        btn_cancelar.pack(side="left", padx=10)


    def cargar_datos_equipo(self, id_equipo):
        """Carga los datos de un equipo existente en el formulario."""
        equipo = database.obtener_equipo_por_id(id_equipo)
        if equipo:
            # Mapeo de índices de la tupla 'equipo' a las claves del diccionario 'entries'
            # (id, asignado_a, tipo_equipo, num_inventario, marca, modelo, num_serie, estado, descripcion_estado, ultima_actualizacion)
            #   0      1           2              3           4       5        6         7             8                     9
            
            for key, widget in self.entries.items():
                if isinstance(widget, ttk.Entry):
                    widget.delete(0, tk.END)
                elif isinstance(widget, ttk.Combobox):
                    widget.set('')
                elif isinstance(widget, tk.Text):
                    widget.delete('1.0', tk.END)

            # Insertar/Establecer valores
            self.entries["asignado_a"].insert(0, equipo[1])
            self.entries["tipo_equipo"].set(equipo[2])
            self.entries["num_inventario"].insert(0, equipo[3])
            self.entries["marca"].insert(0, equipo[4] if equipo[4] else "")
            self.entries["modelo"].insert(0, equipo[5] if equipo[5] else "")
            self.entries["num_serie"].insert(0, equipo[6] if equipo[6] else "")
            self.entries["estado"].set(equipo[7])
            self.entries["descripcion_estado"].insert('1.0', equipo[8] if equipo[8] else "")

    def guardar_equipo(self, id_equipo):
        """Recoge los datos del formulario y los guarda (agrega o actualiza) en la BD."""
        datos = []
        try:
            desc_estado_widget = self.entries["descripcion_estado"]
            desc_estado_val = desc_estado_widget.get('1.0', tk.END).strip()
            
            datos = [
                self.entries["asignado_a"].get(),
                self.entries["tipo_equipo"].get(),
                self.entries["num_inventario"].get(),
                self.entries["marca"].get(),
                self.entries["modelo"].get(),
                self.entries["num_serie"].get(),
                self.entries["estado"].get(),
                desc_estado_val
            ]

            campos_obligatorios_indices = [0, 1, 2, 6]
            campos_obligatorios_nombres = ["Asignado a", "Tipo de Equipo", "N° Inventario/Arrendado", "Estado"]
            
            campos_faltantes = []
            for i, nombre in zip(campos_obligatorios_indices, campos_obligatorios_nombres):
                if not datos[i]:
                    campos_faltantes.append(nombre)

            if campos_faltantes:
                messagebox.showerror("Error de Validación", f"Los siguientes campos son obligatorios:\n- {', '.join(campos_faltantes)}")
                return
                
        except Exception as e:
             messagebox.showerror("Error", f"Error al obtener los datos del formulario: {e}")
             return


        # Lógica para actualizar o agregar
        if id_equipo:
            try:
                if database.actualizar_equipo(id_equipo, datos):
                    messagebox.showinfo("Éxito", "Equipo actualizado correctamente.")
                    self.mostrar_inventario_completo()
                else:
                    messagebox.showerror("Error", "No se pudo actualizar el equipo.\nVerifique que el N° de Serie no esté duplicado.")
            except Exception as e:
                 messagebox.showerror("Error de Base de Datos", f"Ocurrió un error al actualizar: {e}")

        else:
            try:
                if database.agregar_equipo(datos):
                    messagebox.showinfo("Éxito", "Equipo agregado correctamente.")
                    self.mostrar_pantalla_principal()
                else:
                    messagebox.showerror("Error", "No se pudo agregar el equipo.\nEl N° de Serie ya podría existir.")
            except Exception as e:
                 messagebox.showerror("Error de Base de Datos", f"Ocurrió un error al agregar: {e}")


if __name__ == "__main__":
    app = App()
    try:
        app.iconbitmap("tua31v.ico") 
    except tk.TclError:
        print("Advertencia: No se pudo cargar el archivo de icono 'tua31v.ico'. Asegúrate de que el archivo existe.")
    
    app.mainloop()
    